"""2.8 文档只读预览测试：路径安全、超大文档、解析失败、类型限制。"""
from __future__ import annotations

import io

import pytest
from fastapi.testclient import TestClient

from app.services.vault import Vault


def _make_pdf(text: str = "PDF preview keyword docpreview42") -> bytes:
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(72, 500, text)
    c.save()
    return buf.getvalue()


def test_docx_preview(client: TestClient):
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Word 预览测试内容 docpreview-word")
    doc.add_table(rows=1, cols=2)
    doc.tables[0].cell(0, 0).text = "表格内容"
    doc.save(buf)

    from tests.test_rag import _import

    _import(client, "docs", "报告.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    r = client.get("/api/v1/documents/preview", params={"path": "docs/报告.docx"})
    assert r.status_code == 200, r.text
    data = r.json()
    assert "docpreview-word" in data["text"]
    assert "表格内容" in data["text"]
    assert data["name"] == "报告.docx"
    assert data["size"] > 0
    assert data["truncated"] is False


def test_pdf_preview(client: TestClient):
    from tests.test_rag import _import

    _import(client, "docs", "手册.pdf", _make_pdf(), "application/pdf")
    r = client.get("/api/v1/documents/preview", params={"path": "docs/手册.pdf"})
    assert r.status_code == 200, r.text
    assert "docpreview42" in r.json()["text"]


def test_txt_preview_text_interpolation_safe(client: TestClient):
    """预览文本原样返回（前端用 text interpolation 渲染，不经过 HTML）。"""
    from tests.test_rag import _import

    _import(client, "", "note.txt", "<script>alert(1)</script> & <b>x</b>".encode("utf-8"))
    r = client.get("/api/v1/documents/preview", params={"path": "note.txt"})
    assert r.status_code == 200
    assert "<script>alert(1)</script>" in r.json()["text"]


def test_preview_rejects_markdown(client: TestClient):
    _create_md(client, "note.md", "# x")
    r = client.get("/api/v1/documents/preview", params={"path": "note.md"})
    assert r.status_code == 400
    assert "不支持" in r.text


def test_preview_path_security(client: TestClient):
    for bad in ["/etc/passwd", "../secret.txt", "C:\\x.txt", "a/../../b.txt"]:
        r = client.get("/api/v1/documents/preview", params={"path": bad})
        assert r.status_code in (400, 422), bad

    # 隐藏路径 / 回收站
    vault: Vault = client.app.state.vault
    (vault.root / ".hidden.txt").write_text("x", encoding="utf-8")
    r = client.get("/api/v1/documents/preview", params={"path": ".hidden.txt"})
    assert r.status_code == 400
    assert "隐藏" in r.text

    r = client.get("/api/v1/documents/preview", params={"path": ".trash/x.txt"})
    assert r.status_code == 400


def test_preview_missing_file(client: TestClient):
    r = client.get("/api/v1/documents/preview", params={"path": "不存在.txt"})
    assert r.status_code == 404


def test_preview_oversize_rejected(client: TestClient):
    from tests.test_rag import _import

    _import(client, "", "big.txt", b"x" * (20 * 1024 * 1024 + 1))
    r = client.get("/api/v1/documents/preview", params={"path": "big.txt"})
    assert r.status_code == 400
    assert "过大" in r.text


def test_preview_corrupt_pdf(client: TestClient):
    from tests.test_rag import _import

    _import(client, "", "坏.pdf", b"not a real pdf at all")
    r = client.get("/api/v1/documents/preview", params={"path": "坏.pdf"})
    # 解析失败 → 明确错误（400），不返回空成功
    assert r.status_code == 400
    assert "解析失败" in r.text


def test_preview_truncates_long_text(client: TestClient):
    from tests.test_rag import _import

    _import(client, "", "长.txt", ("行" * 200 + "\n") * 3000)  # 约 60 万字符
    r = client.get("/api/v1/documents/preview", params={"path": "长.txt"})
    assert r.status_code == 200
    data = r.json()
    assert data["truncated"] is True
    assert len(data["text"]) <= 500_000


def _create_md(client: TestClient, path: str, content: str) -> None:
    r = client.post("/api/v1/files", json={"path": path, "type": "file", "initial_content": content})
    assert r.status_code == 200, r.text
