"""文档解析：PDF / Word / TXT / CSV → 纯文本，供 AI 索引。

仅解析为文本（不保留格式），输出进入 chunks 表与 .md 同等参与检索。
解析失败返回空文本，由调用方决定是否入块。
"""
from __future__ import annotations

import io
import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)

# 可通过 /api/v1/import 上传并进入 AI 索引的扩展名（md 由原管线处理）
SUPPORTED_EXT = {".pdf", ".docx", ".txt", ".csv", ".md"}
_DOCUMENT_EXT = SUPPORTED_EXT - {".md"}
_MAX_PARSE_CHARS = 2_000_000  # 单文件解析文本上限，防止超大文档拖垮内存


def is_document(path: str) -> bool:
    return Path(path).suffix.lower() in _DOCUMENT_EXT


def parse_document(name: str, data: bytes) -> str:
    """按扩展名提取文本；失败返回空字符串并记录日志。"""
    ext = Path(name).suffix.lower()
    try:
        if ext == ".pdf":
            return _parse_pdf(data)
        if ext == ".docx":
            return _parse_docx(data)
        if ext in (".txt", ".csv"):
            return _parse_text(data)
        return ""
    except Exception as exc:
        logger.warning("解析文档失败 %s: %s", name, exc)
        return ""


def _parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(text)
    return _limit("\n\n".join(parts))


def _parse_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _limit("\n\n".join(parts))


def _parse_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gbk", "latin-1"):
        try:
            return _limit(data.decode(encoding))
        except UnicodeDecodeError:
            continue
    return ""


def _limit(text: str) -> str:
    if len(text) <= _MAX_PARSE_CHARS:
        return text
    return text[: _MAX_PARSE_CHARS] + "\n…（内容过长已截断）"
